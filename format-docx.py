import json
from docx import Document

def add_section_heading(doc, heading):
    doc.add_heading(heading, level=1)

def add_subheading(doc, subheading):
    doc.add_heading(subheading, level=2)

def add_list_items(doc, items):
    for item in items:
        doc.add_paragraph(item)

def add_key_value_pairs(doc, data):
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")

def generate_resume_document(resume_data):
    doc = Document()

    # Add Personal Details section
    add_section_heading(doc, 'Personal Details')
    personal_info = resume_data['personal_details']
    add_key_value_pairs(doc, personal_info)

    # Add Summary section
    add_section_heading(doc, 'Summary')
    summary = resume_data['summary']
    doc.add_paragraph(summary)

    # Add Experience section
    add_section_heading(doc, 'Experience')
    experience = resume_data['experience']
    for exp in experience:
        add_subheading(doc, exp['title'])
        add_key_value_pairs(doc, {'Company': exp['company'], 'Duration': exp['duration']})
        doc.add_paragraph(exp['description'])

    # Add Roles section
    add_section_heading(doc, 'Roles')
    roles = resume_data['roles']
    add_list_items(doc, roles)

    # Add Industries section
    add_section_heading(doc, 'Industries')
    industries = resume_data['industries']
    add_list_items(doc, industries)

    # Add Methods and Processes section
    add_section_heading(doc, 'Methods and Processes')
    methods = resume_data['methods_and_processes']
    add_list_items(doc, methods)

    # Add Certifications section
    add_section_heading(doc, 'Certifications')
    certifications = resume_data['certifications']
    add_list_items(doc, certifications)

    # Add Projects section
    add_section_heading(doc, 'Projects')
    projects = resume_data['projects']
    for project in projects:
        add_subheading(doc, project['title'])
        add_key_value_pairs(doc, {'Company': project['company'], 'Duration': project['duration'], 'Location': project['location']})
        doc.add_paragraph(project['description'])

    # Add Employers section
    add_section_heading(doc, 'Employers')
    employers = resume_data['employers']
    for employer in employers:
        add_subheading(doc, employer['company'])
        doc.add_paragraph(f"Duration: {employer['duration']}")

    # Add Education section
    add_section_heading(doc, 'Education')
    education = resume_data['education']
    for edu in education:
        add_subheading(doc, edu['title'])
        doc.add_paragraph(f"Institution: {edu['institution']}")
        if 'courses' in edu:
            courses = ", ".join(edu['courses'])
            doc.add_paragraph(f"Courses: {courses}")
        doc.add_paragraph(f"Year: {edu['year']}")

    # Add Languages section
    add_section_heading(doc, 'Languages')
    languages = resume_data['languages']
    for lang in languages:
        doc.add_paragraph(f"{lang['language']}: {lang['proficiency']}")

    # Add Engagements and Publications section
    add_section_heading(doc, 'Engagements and Publications')
    engagements = resume_data['engagements_and_publications']
    add_list_items(doc, engagements)

    return doc

def save_document(doc, filename):
    doc.save(filename)

# Read data from 'resume.json'
with open('resume.json', 'r', encoding='utf-8-sig') as f:
    resume_data = json.load(f)

# Generate the resume document
resume_doc = generate_resume_document(resume_data)

# Save the document to a file
save_document(resume_doc, 'resume.docx')
